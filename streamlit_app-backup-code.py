# streamlit_app.py
import os
import uuid
import sqlite3
import math
import numpy as np
import pandas as pd
import streamlit as st
import plotly.express as px
import joblib
import requests
import pycountry
import folium
from streamlit_folium import st_folium
from sklearn.preprocessing import LabelEncoder, StandardScaler
from sklearn.model_selection import train_test_split
from sklearn.ensemble import RandomForestRegressor
from sklearn.metrics import mean_absolute_error, r2_score
from docx import Document
from docx.shared import Inches
from io import BytesIO
from datetime import datetime, timedelta

# --- CONSTANTS ---
OPEN_SKY_API_URL = "https://opensky-network.org/api/states/all"
OPEN_SKY_USER = os.environ.get("OPEN_SKY_USER", "user_placeholder")
OPEN_SKY_PASS = os.environ.get("OPEN_SKY_PASS", "pass_placeholder")

# -------------------------
# Config / folders / DB
# -------------------------
st.set_page_config(page_title="Flight Delay Prediction Dashboard", page_icon="✈️", layout="wide")
st.title("✈️ Flight Delay Prediction Dashboard")
st.caption("Developed by Nirmal Raj | Flight Delay App © 2025")

os.makedirs("models", exist_ok=True)
os.makedirs("outputs", exist_ok=True)
DB_FILE = "prediction_history.db"

# -------------------------
# Airport coordinates (for nearest airport and distance calc)
# -------------------------
AIRPORT_COORDS = {
    "BLR": (12.9716, 77.5946),
    "DEL": (28.7041, 77.1025),
    "BOM": (19.0760, 72.8777),
    "MAA": (13.0827, 80.2707),
    "COK": (9.9312, 76.2673),
    "HYD": (17.3850, 78.4867),
    "CCU": (22.5726, 88.3639),
    "AMD": (23.0225, 72.5714)
}

# -------------------------
# Load or create historical dataset
# -------------------------
DATA_FILE = "flight_delay_data.csv"
MODEL_FILE = "models/best_model.pkl"
COLUMN_RENAMES = {
    "destination": "dest",
    "sched_dep_hour": "departure_time",
    "sched_arr_hour": "arrival_time",
    "delay_minutes": "delay"
}

if os.path.exists(DATA_FILE):
    df = pd.read_csv(DATA_FILE)
    df.rename(columns={k: v for k, v in COLUMN_RENAMES.items() if k in df.columns}, inplace=True)
else:
    st.warning("⚠️ Data not found. Generating synthetic dataset...")
    np.random.seed(42)
    airlines = ["AirA", "AirB", "AirC", "AirD", "AirE"]
    airports = list(AIRPORT_COORDS.keys())
    weather_conditions = ["Clear", "Rain", "Storm", "Fog", "Snow", "Windy"]
    rows = []
    for i in range(1000):
        airline = np.random.choice(airlines)
        origin = np.random.choice(airports)
        dest = np.random.choice([a for a in airports if a != origin])
        dep_hour = np.random.randint(0, 24)
        arr_hour = (dep_hour + np.random.randint(1, 5)) % 24
        distance = np.random.randint(100, 2000)
        day = np.random.randint(0, 7)
        month = np.random.randint(1, 13)
        weather = np.random.choice(weather_conditions, p=[0.6, 0.2, 0.05, 0.08, 0.01, 0.06])
        # Base delay calculation
        delay = max(0, int(np.random.normal(loc=10 + 0.02 * distance + (5 if weather in ["Storm", "Fog"] else 0), scale=15)))
        if np.random.rand() < 0.02:
            delay += np.random.randint(60, 300)
        rows.append({
            "airline": airline, "origin": origin, "destination": dest,
            "sched_dep_hour": dep_hour, "sched_arr_hour": arr_hour,
            "distance": distance, "day_of_week": day, "month": month,
            "weather": weather, "delay_minutes": delay
        })
    df = pd.DataFrame(rows)
    df.to_csv(DATA_FILE, index=False)
    df.rename(columns=COLUMN_RENAMES, inplace=True)

# -------------------------
# Model train/save
# -------------------------
def train_and_save_model(df_local):
    st.info("💡 Training Random Forest Regressor model...")
    le_airline = LabelEncoder()
    le_origin = LabelEncoder()
    le_dest = LabelEncoder()
    le_weather = LabelEncoder()
    le_route = LabelEncoder()
    
    df_train = df_local.copy()
    df_train["airline_enc"] = le_airline.fit_transform(df_train["airline"])
    df_train["origin_enc"] = le_origin.fit_transform(df_train["origin"])
    df_train["dest_enc"] = le_dest.fit_transform(df_train["dest"])
    df_train["weather_enc"] = le_weather.fit_transform(df_train["weather"])
    # NEW: Create Route Feature
    df_train["route"] = df_train["origin"] + "-" + df_train["dest"]
    df_train["route_enc"] = le_route.fit_transform(df_train["route"])
    
    # NEW FEATURE SET (includes route_enc)
    features = ["airline_enc", "origin_enc", "dest_enc", "route_enc", "departure_time", "arrival_time", "distance", "day_of_week", "month", "weather_enc"]
    target = "delay"
    
    scaler = StandardScaler()
    df_train[["departure_time", "arrival_time", "distance", "day_of_week", "month"]] = scaler.fit_transform(
        df_train[["departure_time", "arrival_time", "distance", "day_of_week", "month"]]
    )
    
    X_train, X_test, y_train, y_test = train_test_split(df_train[features], df_train[target], test_size=0.2, random_state=42)
    
    model = RandomForestRegressor(n_estimators=100, random_state=42, n_jobs=-1).fit(X_train, y_train)
    
    y_pred = model.predict(X_test)
    mae = mean_absolute_error(y_test, y_pred)
    r2 = r2_score(y_test, y_pred)
    
    # Save feature importances for XAI
    feature_importances = pd.Series(model.feature_importances_, index=features).sort_values(ascending=False)
    feature_importances.to_csv("outputs/feature_importances.csv", header=True)
    pd.DataFrame({"Metric": ["R²", "MAE"], "Value": [f"{r2:.3f}", f"{mae:.2f}"]}).to_csv("outputs/model_performance.csv", index=False)
    
    joblib.dump(model, MODEL_FILE)
    joblib.dump(le_airline, "models/le_airline.pkl")
    joblib.dump(le_origin, "models/le_origin.pkl")
    joblib.dump(le_dest, "models/le_dest.pkl")
    joblib.dump(le_weather, "models/le_weather.pkl")
    joblib.dump(le_route, "models/le_route.pkl") 
    joblib.dump(scaler, "models/scaler.pkl")
    st.success("✅ Model (Random Forest) training complete and saved to models/")

if not os.path.exists(MODEL_FILE):
    train_and_save_model(df.copy())

# -------------------------
# Load Model / Encoders / Scaler
# -------------------------
try:
    model = joblib.load(MODEL_FILE)
    le_airline = joblib.load("models/le_airline.pkl")
    le_origin = joblib.load("models/le_origin.pkl")
    le_dest = joblib.load("models/le_dest.pkl")
    le_weather = joblib.load("models/le_weather.pkl")
    le_route = joblib.load("models/le_route.pkl")
    scaler = joblib.load("models/scaler.pkl")
    feature_importances = pd.read_csv("outputs/feature_importances.csv", index_col=0, header=None)
    
    # FIX: Ensure index (feature names) is string type for .replace()
    feature_importances.index = feature_importances.index.astype(str)
    
    MODEL_LOADED = True
except FileNotFoundError:
    st.error("🚨 Error: Model components not found. Please ensure the synthetic data was generated and the model training ran successfully.")
    MODEL_LOADED = False
    model, le_airline, le_origin, le_dest, le_weather, le_route, scaler = None, None, None, None, None, None, None
    feature_importances = pd.DataFrame()

# -------------------------
# SQLite history
# -------------------------
def init_db():
    conn = sqlite3.connect(DB_FILE, check_same_thread=False)
    cur = conn.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            timestamp TEXT,
            flight TEXT,
            origin_country TEXT,
            latitude REAL,
            longitude REAL,
            temp REAL,
            wind REAL,
            precip REAL,
            cloud REAL,
            congestion TEXT,
            predicted_delay INTEGER,
            reason TEXT
        )
    """)
    conn.commit()
    return conn

DB_CONN = init_db()

def save_history(record):
    cur = DB_CONN.cursor()
    cur.execute("""
        INSERT INTO history (timestamp, flight, origin_country, latitude, longitude, temp, wind, precip, cloud, congestion, predicted_delay, reason)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        record["timestamp"], record["flight"], record["origin_country"],
        record["latitude"], record["longitude"], record.get("temp"),
        record.get("wind"), record.get("precip"), record.get("cloud"),
        record.get("congestion"), record.get("predicted_delay"),
        record.get("reason")
    ))
    DB_CONN.commit()

def load_history():
    cur = DB_CONN.cursor()
    cur.execute("SELECT * FROM history ORDER BY timestamp DESC")
    columns = [desc[0] for desc in cur.description]
    return pd.DataFrame(cur.fetchall(), columns=columns)

# NEW: Function to clear history
def clear_history():
    cur = DB_CONN.cursor()
    cur.execute("DROP TABLE IF EXISTS history")
    DB_CONN.commit()
    init_db()

# -------------------------
# Helpers
# -------------------------
def haversine(lat1, lon1, lat2, lon2):
    R = 6371.0
    phi1 = math.radians(lat1); phi2 = math.radians(lat2)
    dphi = math.radians(lat2 - lat1); dlambda = math.radians(lon2 - lon1)
    a = math.sin(dphi/2)**2 + math.cos(phi1)*math.cos(phi2)*math.sin(dlambda/2)**2
    return R * 2 * math.atan2(math.sqrt(a), math.sqrt(1-a))

def nearest_airport(lat, lon, airports=AIRPORT_COORDS):
    nearest = None
    min_d = 1e9
    for code, (alat, alon) in airports.items():
        d = haversine(lat, lon, alat, alon)
        if d < min_d: min_d = d; nearest = (code, alat, alon, d)
    return nearest

def compute_congestion(df_live, lat, lon, radius_km=100):
    if df_live is None or df_live.empty: return 0, "Unknown"
    if "latitude" not in df_live.columns or "longitude" not in df_live.columns: return 0, "Unknown"
    counts = 0
    for _, r in df_live.iterrows():
        try:
            la = float(r["latitude"]); lo = float(r["longitude"])
            if np.isnan(la) or np.isnan(lo): continue
            d = haversine(lat, lon, la, lo)
            if d <= radius_km: counts += 1
        except: continue
    if counts < 10: level = "Low"
    elif counts < 25: level = "Medium"
    else: level = "High"
    return counts, level

def safe_encode(le, value):
    """Safely encode a categorical value, returning a default (0) if unseen."""
    try:
        if value in le.classes_:
            return le.transform([value])[0]
        else:
            return 0
    except (ValueError, AttributeError):
        return 0

# -------------------------
# OpenSky API Placeholder (with Caching)
# -------------------------
@st.cache_data(ttl=7200) # Cache data for 7200 seconds (120 minutes)
def fetch_live_flights(bbox=None): 
    """Placeholder for OpenSky API call. Returns a synthetic live dataset.
    The TTL ensures the data is only re-fetched every 120 minutes."""
    
    if bbox is None:
        bbox = [5.0, 68.0, 35.0, 98.0]

    # Use a fixed seed for consistency within the cache window
    np.random.seed(42) 
    
    states = []
    num_flights = 50 + np.random.randint(-10, 10)
    
    airlines_list = list(le_airline.classes_) if MODEL_LOADED else ["AirA", "AirB"]
    airports_list = list(AIRPORT_COORDS.keys())
    weather_conditions_list = list(le_weather.classes_) if MODEL_LOADED else ["Clear", "Rain", "Storm", "Fog", "Snow", "Windy"]
    
    # --- IMPROVED WEATHER SIMULATION (Weighted Probability) ---
    p_values = [0.65, 0.2, 0.05, 0.05, 0.01, 0.04] 

    for i in range(num_flights):
        lat = np.random.uniform(bbox[0], bbox[2])
        lon = np.random.uniform(bbox[1], bbox[3])
        
        origin_sim = np.random.choice(airports_list)
        dest_sim = np.random.choice([a for a in airports_list if a != origin_sim])
        weather_sim = np.random.choice(weather_conditions_list, p=p_values)
        
        # Ensure first flight is always a known route for easy check
        if i == 0: 
            origin_sim = "DEL" 
            dest_sim = "BOM"
            lat = 20.0 
            lon = 75.0

        state_vector = [
            str(uuid.uuid4())[:8], # icao24
            np.random.choice(airlines_list), # callsign (airline)
            'India', # origin_country
            lon, # lon
            lat, # lat
            origin_sim, # simulated origin airport
            dest_sim, # simulated destination airport
            weather_sim, # weather
        ]
        state_vector.extend([None] * (24 - len(state_vector))) 
        states.append(state_vector)

    df_live = pd.DataFrame(states, columns=[
        'icao24', 'callsign', 'origin_country', 'longitude', 'latitude', 
        'origin_airport', 'destination_airport', 'weather', 
        'last_contact', 'longitude_prev', 'latitude_prev', 'velocity', 'true_track', 
        'vertical_rate', 'sensors', 'baro_altitude', 'squawk', 'geo_altitude', 
        'on_ground', 'spi', 'alert', 'tisb', 'vertical_rate_geo', 
        'vertical_rate_baro' 
    ])
    
    df_live = df_live.dropna(subset=['latitude', 'longitude'])
    return df_live

# -------------------------
# Prediction Function for Live Data
# -------------------------
def predict_live_delay(df_live_local, row, time_offset=0):
    """Predicts delay for a single live flight row."""
    if not MODEL_LOADED: return 0, "Model Error", 0, "Unknown", "N/A"

    # 1. Feature Engineering
    lat, lon = row['latitude'], row['longitude']
    
    origin_airport_code = row['origin_airport'] 
    dest_airport_code = row['destination_airport'] 
    airline_code = row['callsign'] 
    weather_condition = row['weather']
    route_code = origin_airport_code + "-" + dest_airport_code

    origin_coords = AIRPORT_COORDS.get(origin_airport_code, (0,0))
    dest_coords = AIRPORT_COORDS.get(dest_airport_code, (0,0))
    distance_km = haversine(origin_coords[0], origin_coords[1], dest_coords[0], dest_coords[1])
    distance_km = max(100, distance_km) 
    
    # Time Feature Engineering (Updated with time offset)
    current_dt = datetime.now() + timedelta(hours=time_offset)
    dep_hour = current_dt.hour
    day_of_week = current_dt.weekday()
    month = current_dt.month
    estimated_flight_duration_hours = max(1, int(distance_km / 800))
    arr_hour = (dep_hour + estimated_flight_duration_hours) % 24
    
    dep_hour = int(dep_hour)
    arr_hour = int(arr_hour)
    day_of_week = int(day_of_week)
    month = int(month)
    # End Time Feature Engineering

    congestion_count, congestion_level = compute_congestion(df_live_local, lat, lon)
    
    # 2. Encoding 
    airline_enc = safe_encode(le_airline, airline_code)
    origin_enc = safe_encode(le_origin, origin_airport_code)
    dest_enc = safe_encode(le_dest, dest_airport_code) 
    weather_enc = safe_encode(le_weather, weather_condition)
    route_enc = safe_encode(le_route, route_code)

    # 3. Scaling Numeric Features
    features_raw = np.array([
        dep_hour, arr_hour, distance_km, day_of_week, month
    ]).reshape(1, -1)
    
    try:
        features_scaled = scaler.transform(features_raw)
    except Exception:
        features_scaled = features_raw 
    
    # 4. Final Feature Vector 
    X = np.array([
        airline_enc, origin_enc, dest_enc, route_enc,
        features_scaled[0, 0], features_scaled[0, 1], features_scaled[0, 2], 
        features_scaled[0, 3], features_scaled[0, 4], 
        weather_enc
    ]).reshape(1, -1)
    
    # 5. Prediction
    prediction = model.predict(X)[0]
    
    predicted_delay = max(0, int(prediction))

    # 6. Reason extraction & Top Factor identification
    top_feature_name = "N/A"
    if not feature_importances.empty:
        # This is safe due to the string conversion fix on load
        top_feature_name_raw = feature_importances.index[0] 
        top_feature_name = top_feature_name_raw.replace('_', ' ').title().replace('Enc', '') # Clean up for display
        
    if weather_condition in ["Storm", "Fog"]:
        reason = f"Weather: {weather_condition}"
    elif congestion_level == "High":
        reason = f"Congestion: {congestion_level} ({congestion_count} flights)"
    elif predicted_delay > 60:
        reason = f"Major Delay Potential (High Factor: {top_feature_name})"
    else:
        reason = "Minor Operational Delay"
        
    return predicted_delay, reason, congestion_count, congestion_level, top_feature_name

# -------------------------
# Tabs
# -------------------------
tabs = st.tabs(["🌍 Live Flights", "📊 Overview", "📈 Analytics", "🤖 Predictions", "📄 Reports"])

# --- 1. Live Flights (tabs[0]) ---
with tabs[0]:
    st.header("🌍 Live Flight Tracking & Delay Prediction")
    
    if not MODEL_LOADED:
        st.error("Cannot run live prediction because the model failed to load.")
    else:
        # --- Fetching data (The function is now cached!) ---
        df_live_base = fetch_live_flights()
        df_live = df_live_base.copy() # Use a mutable copy for filtering
        
        # 1. Filter Controls 
        st.sidebar.header("Live Map Filters")
        delay_filter = st.sidebar.slider("Min Predicted Delay (min)", 0, 90, 5) 
        time_offset = st.sidebar.slider("Simulated Future Time Offset (hours)", 0, 8, 0)
        
        # --- NEW DROPDOWNS ---
        # 1. Country Filter
        country_options = ['All'] + sorted(df_live_base['origin_country'].unique().tolist())
        selected_country = st.sidebar.selectbox("Filter by Origin Country", country_options)

        # 2. Flight ID Filter
        flight_options = ['All'] + sorted(df_live_base['icao24'].unique().tolist())
        selected_flight = st.sidebar.selectbox("Filter by Flight ID", flight_options)
        # ---------------------
        
        col1, col2 = st.columns([3, 1])
        
        # 2. Apply Filters to df_live
        if selected_country != 'All':
            df_live = df_live[df_live['origin_country'] == selected_country]
        if selected_flight != 'All':
            df_live = df_live[df_live['icao24'] == selected_flight]


        with col2:
            st.markdown("### Controls")
            center_lat, center_lon = 20.5937, 78.9629
            
            total_flights_filtered = len(df_live)
            st.metric(label="Total Live Flights Tracked (Simulated)", value=total_flights_filtered)
            
            st.write(f"Showing flights with delay $\ge$ **{delay_filter} min**.")
            st.caption(f"Map time: {datetime.now() + timedelta(hours=time_offset):%H:%M}")
            
            m = folium.Map(location=[center_lat, center_lon], zoom_start=4)
            
            map_data = []
            
            for index, row in df_live.iterrows():
                try:
                    # UPDATED CALL: Pass time_offset and expect top_factor
                    predicted_delay, reason, congestion_count, congestion_level, top_factor = predict_live_delay(df_live, row, time_offset)
                    
                    if predicted_delay < delay_filter:
                        continue # Apply filter
                
                    save_history({
                        "timestamp": datetime.now().isoformat(),
                        "flight": row['icao24'],
                        "origin_country": row['origin_country'],
                        "latitude": row['latitude'],
                        "longitude": row['longitude'],
                        "temp": None, 
                        "wind": None,
                        "precip": None,
                        "cloud": None,
                        "congestion": congestion_level,
                        "predicted_delay": predicted_delay,
                        "reason": reason
                    })
                    
                    if predicted_delay < 15: color = 'green'
                    elif predicted_delay < 45: color = 'orange'
                    else: color = 'red'

                    # Add flight position marker
                    folium.Marker(
                        location=[row['latitude'], row['longitude']],
                        popup=folium.Popup(f"**Flight:** {row['icao24']}<br><b>Delay:</b> {predicted_delay} min<br><b>Reason:</b> {reason}", max_width=300),
                        icon=folium.Icon(color=color, icon='plane', prefix='fa')
                    ).add_to(m)

                    # Add Polyline for route visualization
                    dest_coords = AIRPORT_COORDS.get(row['destination_airport'])
                    if dest_coords:
                        folium.PolyLine(
                            locations=[[row['latitude'], row['longitude']], dest_coords],
                            color='blue',
                            weight=2,
                            opacity=0.5,
                            popup=f"Route: {row['origin_airport']} to {row['destination_airport']}"
                        ).add_to(m)


                    # UPDATED map_data: Added Top Factor
                    map_data.append({
                        "icao24": row['icao24'],
                        "Route": f"{row['origin_airport']}→{row['destination_airport']}",
                        "Delay (min)": predicted_delay,
                        "Reason": reason,
                        "Top Factor": top_factor
                    })
                except Exception as e:
                    st.error(f"FATAL LOOP ERROR: {e}") 
                    continue
            
            # Final Success message 
            st.success(f"✅ Displaying {len(map_data)} filtered flights. History saved.")
            
        with col1:
            st_folium(m, height=500, width='100%', returned_objects=[])

        st.markdown("### Real-time Prediction Table (Filtered)")
        df_map = pd.DataFrame(map_data)
        st.dataframe(df_map, height=200, use_container_width=True)

# --- 2. Overview (tabs[1]) ---
with tabs[1]:
    st.header("📊 Historical & Simulated Overview")
    
    df['route'] = df['origin'].str.cat(df['dest'], sep='-')
    
    col1, col2, col3, col4 = st.columns(4)
    avg_delay = df['delay'].mean()
    long_delays = (df['delay'] > 60).sum()
    
    col1.metric("Total Historical Flights", f"{len(df):,}")
    col2.metric("Average Delay (min)", f"{avg_delay:.2f}")
    col3.metric("Flights Delayed > 60 min", f"{long_delays:,}")
    col4.metric("Most Frequent Route", df['route'].mode()[0])
    
    st.markdown("---")
    
    st.subheader("Simulated 7-Day Delay Forecast (Historical Baseline)")
    
    np.random.seed(123)
    forecast_dates = [datetime.now().date() + timedelta(days=i) for i in range(7)]
    forecast_delay = [max(5, int(np.random.normal(avg_delay, 5) + (3 if d.weekday() in [4, 5, 6] else 0))) for d in forecast_dates]
    forecast_df = pd.DataFrame({
        "Date": forecast_dates,
        "Day": [d.strftime('%a') for d in forecast_dates],
        "Predicted Avg. Delay (min)": forecast_delay
    })
    
    fig_forecast = px.bar(forecast_df, x="Day", y="Predicted Avg. Delay (min)", 
                          title="Average Predicted Daily Delay (Historical Baseline)",
                          color="Predicted Avg. Delay (min)",
                          color_continuous_scale=px.colors.sequential.Sunset,
                          height=400)
    st.plotly_chart(fig_forecast, use_container_width=True)

# --- 3. Analytics (tabs[2]) ---
with tabs[2]:
    st.header("📈 Deep Dive Analytics")
    
    # 1. Model Feature Importance (XAI)
    if not feature_importances.empty:
        st.subheader("1. Model Feature Importance (Random Forest)")
        fig_importance = px.bar(feature_importances, 
                                x=feature_importances.index.str.replace('_', ' ').str.title().str.replace('Enc', ''), 
                                y=feature_importances.iloc[:, 0], 
                                title="Feature Importance Ranking (Gini)",
                                labels={'x': 'Feature', 'y': 'Importance Score'})
        st.plotly_chart(fig_importance, use_container_width=True)
    
    st.markdown("---")
    
    # 2. Historical Delay Distribution
    st.subheader(f"2. Historical Delay Distribution (Flights up to 120 Minutes)")
    delay_data_filtered = df[df['delay'] < 120]
    fig_hist = px.histogram(delay_data_filtered, x="delay", 
                            title=f"Delay Distribution for {len(delay_data_filtered):,} Flights", 
                            nbins=40)
    st.plotly_chart(fig_hist, use_container_width=True)
    
    st.markdown("---")

    # 3. Heatmap
    st.subheader("3. Average Delay Heatmap: Time of Day vs. Day of Week")
    
    # Prep data for Heatmap
    df_heatmap = df.groupby(['day_of_week', 'departure_time'])['delay'].mean().reset_index()
    df_heatmap.columns = ['Day of Week', 'Hour of Day', 'Average Delay (min)']
    
    day_names = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
    df_heatmap['Day of Week Name'] = df_heatmap['Day of Week'].apply(lambda x: day_names[x])

    fig_heatmap = px.density_heatmap(df_heatmap, 
                                     x="Hour of Day", 
                                     y="Day of Week Name", 
                                     z="Average Delay (min)", 
                                     category_orders={"Day of Week Name": day_names[::-1], "Hour of Day": list(range(24))},
                                     title="Delay by Hour and Day (Darker = Longer Delay)",
                                     color_continuous_scale="Reds")
    st.plotly_chart(fig_heatmap, use_container_width=True)

    st.markdown("---")
    
    st.subheader("Average Delay by Categorical Factor")
    
    col1, col2 = st.columns(2)
    
    with col1:
        # 4. Average Delay by Airline
        avg_delay_airline = df.groupby('airline')['delay'].mean().sort_values(ascending=False).reset_index()
        fig_airline = px.bar(avg_delay_airline, x='airline', y='delay', 
                             title="4. Average Delay by Airline", 
                             labels={'delay': 'Avg. Delay (min)', 'airline': 'Airline'})
        st.plotly_chart(fig_airline, use_container_width=True)
    
    with col2:
        # 5. Average Delay by Weather Condition
        avg_delay_weather = df.groupby('weather')['delay'].mean().sort_values(ascending=False).reset_index()
        fig_weather = px.bar(avg_delay_weather, x='weather', y='delay', 
                             title="5. Average Delay by Weather Condition",
                             labels={'delay': 'Avg. Delay (min)', 'weather': 'Weather'})
        st.plotly_chart(fig_weather, use_container_width=True)

# --- 4. Predictions (tabs[3]) ---
with tabs[3]:
    st.header("🤖 Manual Flight Delay Prediction")
    
    if not MODEL_LOADED:
        st.error("Prediction service is unavailable. Model components could not be loaded.")
    else:
        with st.form("prediction_form"):
            col1, col2, col3 = st.columns(3)
            
            with col1:
                input_airline = st.selectbox("Airline", options=le_airline.classes_, index=0)
                input_origin = st.selectbox("Origin Airport", options=le_origin.classes_, index=0)
                input_dest = st.selectbox("Destination Airport", options=le_dest.classes_, index=1)
            
            with col2:
                # UX IMPROVEMENT: Clearer time input
                dep_time_24 = st.slider("Scheduled Departure Hour (24-Hr)", 0, 23, 10)
                arr_time_24 = st.slider("Scheduled Arrival Hour (24-Hr)", 0, 23, 12)
                
                input_dep_time = dep_time_24
                input_arr_time = arr_time_24
                
                # Display 12-hour format for better UX
                dep_ampm = 'AM' if dep_time_24 < 12 else 'PM'
                arr_ampm = 'AM' if arr_time_24 < 12 else 'PM'
                st.caption(f"Departure: {dep_time_24%12 if dep_time_24%12 != 0 else 12} {dep_ampm} | Arrival: {arr_time_24%12 if arr_time_24%12 != 0 else 12} {arr_ampm}")

                input_distance = st.number_input("Flight Distance (km)", 100, 3000, 1000)
            
            with col3:
                input_day = st.selectbox("Day of Week (0=Mon, 6=Sun)", options=list(range(7)), format_func=lambda x: ['Mon', 'Tue', 'Wed', 'Thu', 'Fri', 'Sat', 'Sun'][x])
                input_month = st.selectbox("Month (1-12)", options=list(range(1, 13)))
                input_weather = st.selectbox("Weather Condition", options=le_weather.classes_, index=0)
            
            submitted = st.form_submit_button("Predict Delay")

        if submitted:
            # 1. Encoding
            airline_enc = safe_encode(le_airline, input_airline)
            origin_enc = safe_encode(le_origin, input_origin)
            dest_enc = safe_encode(le_dest, input_dest)
            weather_enc = safe_encode(le_weather, input_weather)
            route_code = input_origin + "-" + input_dest
            route_enc = safe_encode(le_route, route_code) 
            
            # 2. Scaling
            features_raw = np.array([
                input_dep_time, input_arr_time, input_distance, input_day, input_month
            ]).reshape(1, -1)
            
            features_scaled = scaler.transform(features_raw)
            
            # 3. Final Feature Vector 
            X_input = np.array([
                airline_enc, origin_enc, dest_enc, route_enc,
                features_scaled[0, 0], features_scaled[0, 1], features_scaled[0, 2], 
                features_scaled[0, 3], features_scaled[0, 4], 
                weather_enc
            ]).reshape(1, -1)
            
            # 4. Prediction
            prediction = model.predict(X_input)[0]
            
            predicted_delay = max(0, int(prediction))
            
            st.markdown("### Prediction Result")
            if predicted_delay < 15:
                st.success(f"Minimal Delay Expected: **{predicted_delay} minutes**.")
            elif predicted_delay < 45:
                st.warning(f"Moderate Delay Expected: **{predicted_delay} minutes**. Plan accordingly.")
            else:
                st.error(f"Significant Delay Expected: **{predicted_delay} minutes**. Consider re-scheduling.")
            
            r2_score_val = pd.read_csv('outputs/model_performance.csv').iloc[0, 1]
            st.markdown(f"**Details:** Prediction based on Random Forest model (R²: {r2_score_val})")

# --- 5. Reports (tabs[4]) ---
def create_delay_report(df_hist, df_metrics):
    """Generates a DOCX report."""
    doc = Document()
    doc.add_heading('Flight Delay Prediction Report', 0)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("---")
    
    doc.add_heading('1. Model Performance Summary (Random Forest)', level=1)
    doc.add_paragraph(f"R² Score: {df_metrics.iloc[0, 1]}")
    doc.add_paragraph(f"Mean Absolute Error (MAE): {df_metrics.iloc[1, 1]}")

    doc.add_heading('2. Live Prediction History', level=1)
    
    if df_hist.empty:
        doc.add_paragraph("No live prediction history recorded yet.")
    else:
        top_delays = df_hist.sort_values(by='predicted_delay', ascending=False).head(10)
        
        doc.add_heading('Top 10 Most Delayed Predictions (from History)', level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Flight'
        hdr_cells[1].text = 'Time'
        hdr_cells[2].text = 'Predicted Delay (min)'
        hdr_cells[3].text = 'Reason'
        
        for index, row in top_delays.iterrows():
            cells = table.add_row().cells
            cells[0].text = row['flight']
            cells[1].text = datetime.fromisoformat(row['timestamp']).strftime('%H:%M')
            cells[2].text = str(row['predicted_delay'])
            cells[3].text = row['reason']

    return doc

with tabs[4]:
    st.header("📄 Reports & Data Export")
    
    st.subheader("Prediction History")
    df_history = load_history()
    
    col_hist_1, col_hist_2 = st.columns(2)

    with col_hist_1:
        # NEW: Clear History Button
        if st.button("🔴 Clear Prediction History (DANGER ZONE)"): 
            clear_history()
            st.success("Prediction history cleared successfully. Rerunning app...")
            st.rerun()
    
    if df_history.empty:
        st.info("The live prediction history database is currently empty. Run the 'Live Flights' tab to populate it.")
        st.markdown("---")
    else:
        st.dataframe(df_history, height=250, use_container_width=True)
        
        col1, col2 = st.columns(2)
        
        with col1:
            csv_data = df_history.to_csv(index=False).encode('utf-8')
            st.download_button(
                label="📥 Download History as CSV",
                data=csv_data,
                file_name=f'flight_delay_history_{datetime.now().strftime("%Y%m%d")}.csv',
                mime='text/csv'
            )
            
        with col2:
            model_metrics = pd.read_csv("outputs/model_performance.csv")
            doc = create_delay_report(df_history, model_metrics) 
            
            bio = BytesIO()
            doc.save(bio)
            bio.seek(0)
            
            st.download_button(
                label="📝 Download Report as DOCX",
                data=bio.getvalue(),
                file_name=f'flight_delay_report_{datetime.now().strftime("%Y%m%d")}.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

st.caption("© 2025 Nirmal Raj")